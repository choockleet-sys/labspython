from unittest import result
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
from paket import Calculator 
from paket.devices import Iron, Tv, Washer
import tkinter as tk
class App(tk.Tk):
        def __init__(self):
            super().__init__()
            self.title("Расчет электроэнергии")
            self.geometry("500x500")

            self.last_result = None  
        

            self.calc = Calculator(cost_hours = 6.0)
            self.devices ={
                "Iron": Iron(1200),
                "Tv": Tv(100),
                "Washer":Washer(2000)
                }
            #Выбор прибора
            label = tk.Label(self, text="Расчет потребления электроэнергии", font=("Arial", 16, "bold"))
            label.pack(pady=10)
        
            tk.Label(self, text="Выберите прибор:").pack()
            self.device_var = tk.StringVar(value="Iron")
           
            for device_key in self.devices:
                device = self.devices[device_key]
                tk.Radiobutton(self, text=str(device), variable=self.device_var, 
                              value=device_key).pack()

            #Часы
            tk.Label(self, text="Часы использования").pack(pady=(10,0))
            self.hours_entry = tk.Entry(self, width=20)
            self.hours_entry.pack()
            #Цена
            tk.Label(self, text="Цена за кВТ * ч").pack(pady=(10,0))
            self.price_entry = tk.Entry(self, width=20)
            self.price_entry.insert(0, "6.0")
            self.price_entry.pack()
            #Кнопки
            button_frame = tk.Frame(self)
            button_frame.pack(pady=15)
            #Рассчет
            calc_btn = tk.Button(button_frame, text="Рассчитать", command=self.calculate, bg="green", fg="white", font=("Arial", 12))
            calc_btn.pack(side=tk.LEFT, padx=5)
            #Сохранение
            save_btn = tk.Button(button_frame, text="Сохранить отчет", command=self.save_report, bg="blue", fg="white", font=("Arial", 12))
            save_btn.pack(side=tk.LEFT, padx=5)
            #Результат
            self.result_text = tk.Text(self, height=8, width=50)
            self.result_text.pack(pady=10)

        def calculate(self):
            try:
                hours = float(self.hours_entry.get())
                price = float(self.price_entry.get())
                device_key = self.device_var.get()
                
                if hours <= 0:
                    raise ValueError("Часы должны быть больше 0")
                if price <= 0:
                    raise ValueError("Цена должна быть больше 0")

                device = self.devices[device_key]
                self.calc.cost_hours = price

                energy = self.calc.calc_energy(device.power, hours)
                cost = self.calc.calc_cost(device.power, hours)

                self.last_result = {
                "device": str(device),
                "power": device.power,
                "hours": hours,
                "price": price,
                "energy": energy,
                "cost": cost,
                "date": datetime.now().strftime("%d.%m.%Y %H:%M:%S")
            }
                result = f"""
Прибор: {device}
Мощность: {device.power}Вт
Время использования: {hours} часов

Потребление энергии: {energy:.2f} кВт·ч
Цена за кВт·ч: {price} руб.
Общая стоимость: {cost:.2f} руб.
            """
            
                self.result_text.delete(1.0, tk.END)
                self.result_text.insert(1.0, result)
            except ValueError as e:
                self.result_text.delete(1.0, tk.END)
                self.result_text.insert(1.0, f" Ошибка! {str(e)}")
    


        def save_report(self):
            if not self.last_result:
                messagebox.showwarning("Внимание", "Сначала выполните расчет!")
                return
            file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word документы", "*.docx"), ("Все файлы", "*.*")],
            initialfile=f"Отчет_электроэнергия_{datetime.now().strftime('%d_%m_%Y')}.docx"
        )
            if not file_path:
                return
            try:
                # Создаем документ
                doc = Document()
            
                # Заголовок
                title = doc.add_heading("ОТЧЕТ О РАСЧЕТЕ ЭЛЕКТРОЭНЕРГИИ", 0)
                title.alignment = 1  # Центрирование
            
                # Дата
                date_para = doc.add_paragraph(f"Дата создания: {self.last_result['date']}")
                date_para.alignment = 1
            
                # Основная информация
                doc.add_heading("Информация о приборе", level=1)
            
                table = doc.add_table(rows=5, cols=2)
                table.style = 'Light Grid Accent 1'
            
                # Заполняем таблицу
                table.rows[0].cells[0].text = "Прибор"
                table.rows[0].cells[1].text = self.last_result['device']
            
                table.rows[1].cells[0].text = "Мощность"
                table.rows[1].cells[1].text = f"{self.last_result['power']} Вт"
            
                table.rows[2].cells[0].text = "Время использования"
                table.rows[2].cells[1].text = f"{self.last_result['hours']} часов"
            
                table.rows[3].cells[0].text = "Цена за кВт·ч"
                table.rows[3].cells[1].text = f"{self.last_result['price']} руб."
            
                table.rows[4].cells[0].text = "Потребление энергии"
                table.rows[4].cells[1].text = f"{self.last_result['energy']:.2f} кВт·ч"
            
                # Итоги
                doc.add_heading("Итоговая стоимость", level=1)
            
                cost_para = doc.add_paragraph()
                cost_run = cost_para.add_run(f"{self.last_result['cost']:.2f} руб.")
                cost_run.font.size = Pt(18)
                cost_run.font.bold = True
                cost_run.font.color.rgb = RGBColor(255, 0, 0)
                cost_para.alignment = 1
            
                # Сохраняем файл
                doc.save(file_path)
                messagebox.showinfo("Успех", f"Отчет сохранен в:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")
if __name__ == "__main__":
    app = App()
    app.mainloop()
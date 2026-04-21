# Импортируем необходимые модули Kivy
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.togglebutton import ToggleButton
from kivy.uix.scrollview import ScrollView
from kivy.uix.popup import Popup

# Стандартные библиотеки
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor

# Наши классы приборов и калькулятора
from devices import Iron, Tv, Washer
from calculator import Calculator


# Главный виджет приложения — вся компоновка экрана
class MainLayout(BoxLayout):
    def __init__(self, **kwargs):
        # Последний результат расчёта (нужен для сохранения отчёта)
        self.last_result = None

        # Создаём калькулятор с тарифом по умолчанию
        self.calc = Calculator(cost_per_kwh=6.0)

        # Словарь приборов — создаём ДО вызова super().__init__()
        # иначе Kivy может обратиться к self.devices раньше времени
        self.devices = {
            "Iron":   Iron(1200),
            "Tv":     Tv(100),
            "Washer": Washer(2000)
        }

        # Вызываем конструктор родительского класса после инициализации данных
        super().__init__(**kwargs)

        # Вертикальная ориентация — элементы идут сверху вниз
        self.orientation = "vertical"
        self.padding = 20
        self.spacing = 10

        # Строим интерфейс
        self._build_ui()

    def __str__(self):
        return f"MainLayout(devices={list(self.devices.keys())})"

    def __repr__(self):
        return f"MainLayout(calc={self.calc})"

    def _build_ui(self):
        # Заголовок
        title = Label(
            text="Расчет потребления электроэнергии",
            font_size=20,
            bold=True,
            size_hint_y=None,
            height=50
        )
        self.add_widget(title)

        # Подпись над кнопками выбора прибора
        self.add_widget(Label(
            text="Выберите прибор:",
            size_hint_y=None,
            height=30
        ))

        # Кнопки выбора прибора (ToggleButton — как радиокнопки)
        self.toggle_buttons = {}
        for i, (key, device) in enumerate(self.devices.items()):
            btn = ToggleButton(
                text=str(device),
                group="devices",        # group — только одна кнопка активна
                size_hint_y=None,
                height=40
            )
            # Первый прибор выбран по умолчанию
            if i == 0:
                btn.state = "down"
            self.toggle_buttons[key] = btn
            self.add_widget(btn)

        # Поле ввода часов
        self.add_widget(Label(
            text="Часы использования:",
            size_hint_y=None,
            height=30
        ))
        self.hours_input = TextInput(
            hint_text="например, 2.5",
            multiline=False,
            size_hint_y=None,
            height=40
        )
        self.add_widget(self.hours_input)

        # Поле ввода цены
        self.add_widget(Label(
            text="Цена за кВт·ч (руб.):",
            size_hint_y=None,
            height=30
        ))
        self.price_input = TextInput(
            text="6.0",
            multiline=False,
            size_hint_y=None,
            height=40
        )
        self.add_widget(self.price_input)

        # Кнопки "Рассчитать" и "Сохранить отчёт" рядом
        btn_row = BoxLayout(
            orientation="horizontal",
            size_hint_y=None,
            height=45,
            spacing=10
        )

        calc_btn = Button(
            text="Рассчитать",
            background_color=(0, 0.6, 0, 1)    # зелёный
        )
        calc_btn.bind(on_press=self.calculate)

        save_btn = Button(
            text="Сохранить отчёт",
            background_color=(0, 0, 0.8, 1)    # синий
        )
        save_btn.bind(on_press=self.save_report)

        btn_row.add_widget(calc_btn)
        btn_row.add_widget(save_btn)
        self.add_widget(btn_row)

        # Поле вывода результата (с прокруткой)
        scroll = ScrollView()
        self.result_label = Label(
            text="Результат появится здесь...",
            size_hint_y=None,
            halign="left",
            valign="top"
        )
        # Привязываем ширину текста к ширине виджета
        self.result_label.bind(
            width=lambda *x: self.result_label.setter("text_size")(
                self.result_label, (self.result_label.width, None)
            ),
            texture_size=lambda *x: self.result_label.setter("height")(
                self.result_label, self.result_label.texture_size[1]
            )
        )
        scroll.add_widget(self.result_label)
        self.add_widget(scroll)

    def _get_selected_key(self):
        # Возвращаем ключ выбранного прибора
        for key, btn in self.toggle_buttons.items():
            if btn.state == "down":
                return key
        return "Iron"

    def _show_popup(self, title, message):
        # Вспомогательный метод — показывает всплывающее окно с сообщением
        content = BoxLayout(orientation="vertical", padding=10, spacing=10)
        content.add_widget(Label(text=message))
        close_btn = Button(text="OK", size_hint_y=None, height=40)
        content.add_widget(close_btn)

        popup = Popup(title=title, content=content, size_hint=(0.8, 0.4))
        close_btn.bind(on_press=popup.dismiss)
        popup.open()

    def calculate(self, instance):
        # Обработчик кнопки "Рассчитать"
        try:
            hours = float(self.hours_input.text)
            price = float(self.price_input.text)

            if hours <= 0:
                raise ValueError("Часы должны быть больше 0")
            if price <= 0:
                raise ValueError("Цена должна быть больше 0")

            key = self._get_selected_key()
            device = self.devices[key]

            # Устанавливаем тариф через managed-атрибут
            self.calc.cost_per_kwh = price

            energy = self.calc.calc_energy(device.power, hours)
            cost = self.calc.calc_cost(device.power, hours)

            # Сохраняем результат для отчёта
            self.last_result = {
                "device": str(device),
                "power":  device.power,
                "hours":  hours,
                "price":  price,
                "energy": energy,
                "cost":   cost,
                "date":   datetime.now().strftime("%d.%m.%Y %H:%M:%S")
            }

            # Выводим результат в текстовое поле
            self.result_label.text = (
                f"Прибор: {device}\n"
                f"Мощность: {device.power} Вт\n"
                f"Время: {hours} ч\n\n"
                f"Потребление: {energy:.2f} кВт·ч\n"
                f"Цена за кВт·ч: {price} руб.\n"
                f"Итого: {cost:.2f} руб."
            )

        except ValueError as e:
            self._show_popup("Ошибка", str(e))

    def save_report(self, instance):
        # Обработчик кнопки "Сохранить отчёт"
        if not self.last_result:
            self._show_popup("Внимание", "Сначала выполните расчёт!")
            return

        try:
            # Формируем имя файла с датой
            path = f"Отчет_{datetime.now().strftime('%d_%m_%Y')}.docx"

            doc = Document()

            # Заголовок документа
            title = doc.add_heading("ОТЧЕТ О РАСЧЕТЕ ЭЛЕКТРОЭНЕРГИИ", 0)
            title.alignment = 1

            # Дата
            doc.add_paragraph(f"Дата: {self.last_result['date']}").alignment = 1

            # Таблица с данными
            doc.add_heading("Информация о приборе", level=1)
            table = doc.add_table(rows=5, cols=2)
            table.style = "Light Grid Accent 1"

            table.rows[0].cells[0].text = "Прибор"
            table.rows[0].cells[1].text = self.last_result["device"]
            table.rows[1].cells[0].text = "Мощность"
            table.rows[1].cells[1].text = f"{self.last_result['power']} Вт"
            table.rows[2].cells[0].text = "Время использования"
            table.rows[2].cells[1].text = f"{self.last_result['hours']} ч"
            table.rows[3].cells[0].text = "Цена за кВт·ч"
            table.rows[3].cells[1].text = f"{self.last_result['price']} руб."
            table.rows[4].cells[0].text = "Потребление"
            table.rows[4].cells[1].text = f"{self.last_result['energy']:.2f} кВт·ч"

            # Итоговая стоимость
            doc.add_heading("Итого", level=1)
            p = doc.add_paragraph()
            run = p.add_run(f"{self.last_result['cost']:.2f} руб.")
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            p.alignment = 1

            doc.save(path)
            self._show_popup("Успех", f"Отчёт сохранён:\n{path}")

        except Exception as e:
            self._show_popup("Ошибка", str(e))


# Класс приложения Kivy — точка входа
class ElectricApp(App):
    def build(self):
        # build() возвращает корневой виджет приложения
        return MainLayout()

    def __str__(self):
        return "ElectricApp()"

    def __repr__(self):
        return "ElectricApp(title='Расчет электроэнергии')"


# Запуск приложения
if __name__ == "__main__":
    ElectricApp().run()
